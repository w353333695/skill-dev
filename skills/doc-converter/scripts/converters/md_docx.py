"""Markdown → Word (.docx) 转换器

支持的 Markdown 格式：
- 标题 (h1-h6)
- 加粗、斜体、行内代码、删除线
- 链接、图片
- 无序/有序/任务列表（含嵌套）
- 引用块
- 代码块（带语言标识）
- 表格
- 分割线
- Mermaid 图表（自动渲染为图片）
"""

import re
import tempfile
from pathlib import Path
from .base import BaseConverter, ConvertResult, register


def _render_mermaid_images(mermaid_blocks: list[str], tmp_dir: Path) -> list[Path]:
    """将 mermaid 代码块渲染为图片，返回图片路径列表"""
    from .mermaid_image import render_mermaid_to_image
    images = []
    for i, code in enumerate(mermaid_blocks):
        img_path = tmp_dir / f"mermaid_{i}.png"
        render_mermaid_to_image(code, img_path)
        images.append(img_path)
    return images


# 当前转换的输入文件路径，供 _add_inline_runs 解析相对图片路径用
_CURRENT_INPUT_PATH: Path | None = None


def _resolve_image_path(input_path: Path | None, src: str) -> Path | None:
    """解析 markdown 图片 src 为本地文件路径。

    支持绝对路径与相对路径（相对于 md 文件所在目录）。
    http/https 等 URL 与无法解析的路径返回 None（由调用方回退占位文字）。

    Args:
        input_path: 当前 md 文件路径
        src: 图片地址，如 ./_assets/x.png 或 /abs/x.png
    """
    if not src or src.startswith(("http://", "https://", "data:")):
        return None
    p = Path(src)
    if p.is_absolute():
        return p
    if input_path is None:
        return None
    return (input_path.parent / p).resolve()


# 内联格式解析正则
# 顺序很重要：先匹配更长的模式
_INLINE_PATTERNS = [
    # 图片 ![alt](url)
    (r"!\[([^\]]*)\]\(([^)]+)\)", "image"),
    # 链接 [text](url)
    (r"\[([^\]]+)\]\(([^)]+)\)", "link"),
    # 粗斜体 ***text*** 或 ___text___
    (r"(\*\*\*|___)(.+?)\1", "bold_italic"),
    # 加粗 **text** 或 __text__
    (r"(\*\*|__)(.+?)\1", "bold"),
    # 删除线 ~~text~~
    (r"~~(.+?)~~", "strikethrough"),
    # 行内代码 `text`
    (r"`([^`]+)`", "code"),
    # 斜体 *text* 或 _text_（不匹配已消耗的 ** 或 __）
    (r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|(?<!_)_(?!_)(.+?)(?<!_)_(?!_)", "italic"),
]

# 编译为单一分组正则，用于 split
_INLINE_RE = re.compile(
    r"(!\[[^\]]*\]\([^)]+\)"       # 图片
    r"|\[[^\]]+\]\([^)]+\)"        # 链接
    r"|\*\*\*.+?\*\*\*"            # 粗斜体
    r"|___.+?___"                   # 粗斜体
    r"|\*\*.+?\*\*"                # 加粗
    r"|__.+?__"                    # 加粗
    r"|~~.+?~~"                    # 删除线
    r"|`.+?`"                      # 行内代码
    r"|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)"  # 斜体
    r"|(?<!_)_(?!_).+?(?<!_)_(?!_)"        # 斜体
    r")"
)


def _set_run_font(run, font_name: str, font_size=None):
    """正确设置 run 的中英文字体（含 eastAsia）"""
    from docx.oxml.ns import qn

    run.font.name = font_name
    if font_size is not None:
        run.font.size = font_size
    # 关键：设置东亚字体，否则中文不生效
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(style, font_name: str, font_size=None):
    """正确设置样式的中英文字体（含 eastAsia）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    style.font.name = font_name
    if font_size:
        style.font.size = font_size
    # 设置东亚字体
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_inline_runs(paragraph, text: str, base_font_name: str = "Microsoft YaHei",
                     base_font_size=None):
    """解析内联 Markdown 格式并添加 runs 到段落

    Args:
        paragraph: python-docx 段落对象
        text: 含 Markdown 内联格式的文本
        base_font_name: 基础字体名
        base_font_size: 基础字号 (Pt 对象)
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not text:
        return

    # 分割文本为普通文本和格式化片段
    parts = _INLINE_RE.split(text)

    for part in parts:
        if not part:
            continue

        # 图片 ![alt](url)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", part)
        if m:
            alt_text = m.group(1) or "图片"
            src = m.group(2).strip()
            from docx.shared import Inches as _Inches
            # 行内图片：尝试读取并以小尺寸插入；找不到则回退为占位文字
            img_path = _resolve_image_path(_CURRENT_INPUT_PATH, src)
            if img_path and img_path.exists():
                run = paragraph.add_run()
                run.add_picture(str(img_path), width=_Inches(2))
            else:
                run = paragraph.add_run(f"[{alt_text}]")
                _set_run_font(run, base_font_name, base_font_size)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.italic = True
            continue

        # 链接 [text](url)
        m = re.match(r"^\[([^\]]+)\]\(([^)]+)\)$", part)
        if m:
            link_text = m.group(1)
            url = m.group(2)
            _add_hyperlink(paragraph, link_text, url, base_font_name, base_font_size)
            continue

        # 粗斜体 ***text*** 或 ___text___
        m = re.match(r"^(\*\*\*|___)(.+?)\1$", part)
        if m:
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
            _set_run_font(run, base_font_name, base_font_size)
            continue

        # 加粗 **text** 或 __text__
        m = re.match(r"^(\*\*|__)(.+?)\1$", part)
        if m:
            run = paragraph.add_run(m.group(2))
            run.bold = True
            _set_run_font(run, base_font_name, base_font_size)
            continue

        # 删除线 ~~text~~
        m = re.match(r"^~~(.+?)~~$", part)
        if m:
            run = paragraph.add_run(m.group(1))
            run.font.strike = True
            _set_run_font(run, base_font_name, base_font_size)
            continue

        # 行内代码 `text`
        m = re.match(r"^`(.+?)`$", part)
        if m:
            run = paragraph.add_run(m.group(1))
            _set_run_font(run, "Consolas", Pt(10))
            # 设置背景色（浅灰）
            _set_run_shading(run, "E8E8E8")
            continue

        # 斜体 *text* 或 _text_
        m = re.match(r"^\*(.+?)\*$|^_(.+?)_$", part)
        if m:
            content = m.group(1) or m.group(2)
            run = paragraph.add_run(content)
            run.italic = True
            _set_run_font(run, base_font_name, base_font_size)
            continue

        # 普通文本
        run = paragraph.add_run(part)
        _set_run_font(run, base_font_name, base_font_size)


def _add_hyperlink(paragraph, text: str, url: str, font_name: str, font_size):
    """向段落添加超链接"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    # 创建超链接关系
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # 构建超链接 XML
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # 蓝色下划线样式
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.append(rFonts)

    if font_size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size.pt * 2)))
        rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def _set_run_shading(run, color_hex: str):
    """为 run 设置背景色"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    run._r.get_or_add_rPr().append(shd)


def _set_paragraph_shading(paragraph, color_hex: str):
    """为段落设置背景色"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    paragraph._p.get_or_add_pPr().append(shd)


def _add_left_border(paragraph, color_hex: str = "CCCCCC", width: int = 24):
    """为段落添加左边框（引用块样式）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _get_indent_level(line: str) -> int:
    """获取列表缩进层级（每 2 或 4 空格为一级）"""
    spaces = len(line) - len(line.lstrip())
    # 支持 tab 或 2/4 空格缩进
    if "\t" in line[:spaces]:
        return line[:spaces].count("\t")
    if spaces >= 4:
        return spaces // 4
    if spaces >= 2:
        return spaces // 2
    return 0


@register
class MdToDocx(BaseConverter):
    name = "md-docx"
    source_formats = ["md", "markdown"]
    target_formats = ["docx"]
    description = "Markdown 转 Word 文档 (Mermaid 自动转图片插入)"
    dependencies = ["docx"]  # python-docx, import 名为 docx

    # 标题字号映射：h1 → h6
    HEADING_SIZES = {
        1: 22,  # 一级标题 22pt
        2: 18,  # 二级标题 18pt
        3: 15,  # 三级标题 15pt
        4: 13,  # 四级标题 13pt
        5: 12,  # 五级标题 12pt
        6: 11,  # 六级标题 11pt
    }

    def convert(self, input_path: Path, output_path: Path, **options) -> ConvertResult:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        global _CURRENT_INPUT_PATH
        _CURRENT_INPUT_PATH = input_path
        self._embedded_img_count = 0

        text = input_path.read_text(encoding="utf-8")
        doc = Document()

        # 基础样式设置 — 统一微软雅黑
        style = doc.styles["Normal"]
        _set_style_font(style, "Microsoft YaHei", Pt(11))
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.3

        # 覆盖所有标题样式的字体为微软雅黑，各级字号不同
        for level in range(1, 7):
            heading_style = doc.styles[f"Heading {level}"]
            _set_style_font(heading_style, "Microsoft YaHei", Pt(self.HEADING_SIZES[level]))
            heading_style.font.bold = True

        # 预提取 mermaid 块并渲染为图片
        mermaid_pattern = r"```mermaid\s*\n(.*?)```"
        mermaid_blocks = re.findall(mermaid_pattern, text, re.DOTALL)
        mermaid_images = {}

        if mermaid_blocks:
            try:
                tmp_dir = Path(tempfile.mkdtemp(prefix="mermaid_"))
                images = _render_mermaid_images(mermaid_blocks, tmp_dir)
                for i, img in enumerate(images):
                    mermaid_images[i] = img
            except Exception:
                pass

        mermaid_idx = 0
        lines = text.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # === 标题 ===
            m = re.match(r"^(#{1,6})\s+(.+)$", line)
            if m:
                level = len(m.group(1))
                heading_text = m.group(2).strip()
                # 去除标题中的内联格式标记用于 heading
                heading = doc.add_heading(level=level)
                # 标题不传 font_size，让 run 继承标题样式的字号
                _add_inline_runs(heading, heading_text, base_font_size=None)
                i += 1
                continue

            # === Mermaid 代码块 → 插入图片 ===
            if line.strip().startswith("```mermaid"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```

                if mermaid_idx in mermaid_images and mermaid_images[mermaid_idx].exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(mermaid_images[mermaid_idx]), width=Inches(6))
                else:
                    self._add_code_block(doc, code_lines, "mermaid")

                mermaid_idx += 1
                continue

            # === 普通代码块 ===
            if line.strip().startswith("```"):
                lang = line.strip()[3:].strip()
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```
                self._add_code_block(doc, code_lines, lang)
                continue

            # === 表格 ===
            if "|" in line and i + 1 < len(lines) and re.match(
                    r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
                rows = []
                while i < len(lines) and "|" in lines[i]:
                    raw = lines[i].strip()
                    # 跳过分隔行
                    if re.match(r"^\|[\s\-:|]+\|$", raw):
                        i += 1
                        continue
                    cells = [c.strip() for c in raw.strip("|").split("|")]
                    rows.append(cells)
                    i += 1
                if rows:
                    self._add_table(doc, rows)
                continue

            # === 引用块 ===
            if line.startswith(">"):
                quote_lines = []
                while i < len(lines) and (lines[i].startswith(">") or
                        (lines[i].strip() and quote_lines and not lines[i].startswith("#"))):
                    content = re.sub(r"^>\s?", "", lines[i])
                    quote_lines.append(content)
                    i += 1
                self._add_blockquote(doc, quote_lines)
                continue

            # === 任务列表 ===
            m = re.match(r"^(\s*)[-*+]\s+\[([ xX])\]\s+(.+)$", line)
            if m:
                indent = _get_indent_level(line)
                checked = m.group(2).lower() == "x"
                content = m.group(3)
                prefix = "\u2611 " if checked else "\u2610 "
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(1.27 * (indent + 1))
                _add_inline_runs(p, prefix + content)
                i += 1
                continue

            # === 无序列表 ===
            m = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
            if m:
                indent = _get_indent_level(line)
                content = m.group(2)
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(1.27 * (indent + 1))
                _add_inline_runs(p, content)
                i += 1
                continue

            # === 有序列表 ===
            m = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
            if m:
                indent = _get_indent_level(line)
                content = m.group(2)
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.left_indent = Cm(1.27 * (indent + 1))
                _add_inline_runs(p, content)
                i += 1
                continue

            # === 分割线 ===
            if re.match(r"^\s*(---+|\*\*\*+|___+)\s*$", line.strip()):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                # 用底部边框模拟分割线
                self._add_hr_border(p)
                i += 1
                continue

            # === 空行 ===
            if not line.strip():
                i += 1
                continue

            # === 独立行图片 ![alt](path) ===
            # 单独成行的图片整张插入，而非内联占位文字
            m = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
            if m:
                alt_text = m.group(1) or "图片"
                src = m.group(2).strip()
                self._add_standalone_image(doc, input_path, src, alt_text)
                i += 1
                continue

            # === 普通段落 ===
            p = doc.add_paragraph()
            _add_inline_runs(p, line)
            i += 1

        doc.save(str(output_path))

        msg = "Markdown 已转为 Word 文档"
        parts = []
        embedded = getattr(self, "_embedded_img_count", 0)
        if embedded:
            parts.append(f"{embedded} 张本地图片")
        img_count = len([v for v in mermaid_images.values() if v.exists()])
        if img_count:
            parts.append(f"{img_count} 个 Mermaid 图表")
        if parts:
            msg += "，" + "、".join(parts) + "已插入"
        return ConvertResult(True, output_path=output_path, message=msg)

    def _add_standalone_image(self, doc, input_path: Path, src: str, alt_text: str):
        """插入单独成行的图片，读取本地文件并居中嵌入。

        Args:
            doc: python-docx Document 对象
            input_path: 当前 md 文件路径，用于解析相对图片地址
            src: 图片地址（相对/绝对路径，或 URL）
            alt_text: 图片替代文字，作为图注
        """
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        img_path = _resolve_image_path(input_path, src)
        if img_path and img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            # 宽度限制在 A4 可用宽度内（约 6 寸），高度按原图比例缩放
            run.add_picture(str(img_path), width=Inches(6))
            self._embedded_img_count += 1
            # 图注
            if alt_text and alt_text != "图片":
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(alt_text)
                _set_run_font(cap_run, "Microsoft YaHei", Pt(9))
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                cap_run.italic = True
        else:
            # 图片找不到时回退为占位文字
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[图片缺失: {alt_text}]")
            _set_run_font(run, "Microsoft YaHei", Pt(10))
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True

    def _add_code_block(self, doc, code_lines: list[str], lang: str = ""):
        """添加代码块，带背景色和等宽字体"""
        from docx.shared import Pt, Cm

        for line_text in code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0.5)
            _set_paragraph_shading(p, "F5F5F5")
            run = p.add_run(line_text if line_text else " ")
            _set_run_font(run, "Consolas", Pt(9.5))

    def _add_table(self, doc, rows: list[list[str]]):
        """添加表格，首行加粗"""
        from docx.shared import Pt

        if not rows:
            return
        max_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols, style="Table Grid")

        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx >= max_cols:
                    break
                cell = table.cell(r_idx, c_idx)
                cell.text = ""  # 清空默认段落
                p = cell.paragraphs[0]
                _add_inline_runs(p, cell_text)
                # 首行加粗
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True

    def _add_blockquote(self, doc, quote_lines: list[str]):
        """添加引用块，带左边框和灰色背景"""
        from docx.shared import Pt, Cm, RGBColor

        for line_text in quote_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_left_border(p, "4472C4", 18)
            _set_paragraph_shading(p, "F2F7FC")
            if line_text.strip():
                _add_inline_runs(p, line_text)
            else:
                p.add_run(" ")  # 空行占位

    def _add_hr_border(self, paragraph):
        """用底部边框模拟水平分割线"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)
