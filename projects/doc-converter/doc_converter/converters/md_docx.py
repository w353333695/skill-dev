"""Markdown → Word (.docx) 转换器

支持的 Markdown 格式：
- 标题 (h1-h6)
- 加粗、斜体、行内代码、删除线、粗斜体嵌套（内联解析走 markdown 库，正确支持嵌套）
- 链接（行内 / 引用式 / 自动 / 邮箱）、图片
- 无序/有序/任务列表（含嵌套）；有序列表按章节自动重置编号
- 引用块（多层嵌套、引用内代码块/列表/内联）
- 代码块（Pygments 语法高亮，带语言标识）
- 表格（含列对齐 :--: / --: / :--）
- 分割线
- 转义字符（\\* \\_ \\` … 原样显示）
- 脚注（[^id] 上标 + 文末列表）
- 数学公式（$...$ / $$...$$ 防误解析，等宽斜体渲染）
- Mermaid 图表（自动渲染为图片）
- 软换行（行尾两空格 → 段内换行）
"""

import html as _html_lib
import re
import tempfile
from html.parser import HTMLParser
from pathlib import Path
from .base import BaseConverter, ConvertResult, register


# ---------------------------------------------------------------------------
# 内联渲染上下文（模块级，沿用既有全局模式；convert 入口设置）
# ---------------------------------------------------------------------------
_CURRENT_INPUT_PATH: Path | None = None
_MATH_CACHE: list[str] = []   # 公式原文列表（按占位符序号索引）
_FN_CACHE: list[tuple[int, str, str]] = []  # [(序号, id, 内容)]


# ---------------------------------------------------------------------------
# 路径与字体辅助
# ---------------------------------------------------------------------------
def _resolve_image_path(input_path: Path | None, src: str) -> Path | None:
    """解析 markdown 图片 src 为本地文件路径（支持绝对/相对路径）。

    http(s)/data: 等 URL 或无法解析时返回 None（调用方回退占位文字）。
    """
    if not src or src.startswith(("http://", "https://", "data:", "mailto:")):
        return None
    p = Path(src)
    if p.is_absolute():
        return p
    if input_path is None:
        return None
    return (input_path.parent / p).resolve()


def _set_run_font(run, font_name: str, font_size=None):
    """正确设置 run 的中英文字体（含 eastAsia）。"""
    from docx.oxml.ns import qn

    run.font.name = font_name
    if font_size is not None:
        run.font.size = font_size
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(style, font_name: str, font_size=None):
    """正确设置样式的中英文字体（含 eastAsia）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    style.font.name = font_name
    if font_size:
        style.font.size = font_size
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_hyperlink(paragraph, text: str, url: str, font_name: str, font_size):
    """向段落添加超链接（蓝色下划线，含中英文字体）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.append(rFonts)

    if font_size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size.pt * 2)))
        rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def _set_run_shading(run, color_hex: str):
    """为 run 设置背景色。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    run._r.get_or_add_rPr().append(shd)


def _set_paragraph_shading(paragraph, color_hex: str):
    """为段落设置背景色。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    paragraph._p.get_or_add_pPr().append(shd)


def _add_left_border(paragraph, color_hex: str = "CCCCCC", width: int = 24):
    """为段落添加左边框（引用块样式）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _get_indent_level(line: str) -> int:
    """获取列表缩进层级（每 2 或 4 空格为一级，tab 算一级）。"""
    stripped = line.lstrip()
    spaces = len(line) - len(stripped)
    if "\t" in line[:spaces]:
        return line[:spaces].count("\t") + (spaces - line[:spaces].count("\t") * 1) // 4
    if spaces >= 4:
        return spaces // 4
    if spaces >= 2:
        return spaces // 2
    return 0


def _render_mermaid_images(mermaid_blocks: list[str], tmp_dir: Path) -> list[Path]:
    """将 mermaid 代码块渲染为图片，返回图片路径列表。"""
    from .mermaid_image import render_mermaid_to_image
    images = []
    for i, code in enumerate(mermaid_blocks):
        img_path = tmp_dir / f"mermaid_{i}.png"
        render_mermaid_to_image(code, img_path)
        images.append(img_path)
    return images


# ---------------------------------------------------------------------------
# 预处理：公式 / 脚注 / 引用式链接
# ---------------------------------------------------------------------------
# 占位符使用 \x00 包裹，markdown 库与正则都不会触碰
_MATH_BLOCK_RE = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
_MATH_INLINE_RE = re.compile(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)")
_FN_DEF_RE = re.compile(r"^[ \t]{0,3}\[\^([^\]]+)\]:[ \t]*(.+)$", re.MULTILINE)
_FN_REF_RE = re.compile(r"\[\^([^\]]+)\]")
_REF_DEF_RE = re.compile(
    r'^[ \t]{0,3}\[([^\]^]+)\]:[ \t]*(\S+)(?:[ \t]+"([^"]*)")?[ \t]*$',
    re.MULTILINE,
)


def _extract_math(md_text: str) -> tuple[str, list[str]]:
    """把 $$...$$ 与 $...$ 公式替换为占位符，避免其中的 _ * ^ 触发 markdown 格式。

    返回 (处理后的文本, 公式原文列表)。占位符形如 \\x00MATH0B\\x00（块级）/ \\x00MATH0I\\x00（行内）。
    """
    maths: list[str] = []

    def block_repl(m):
        maths.append(m.group(1).strip())
        return f"\x00MATH{len(maths) - 1}B\x00"

    md_text = _MATH_BLOCK_RE.sub(block_repl, md_text)

    def inline_repl(m):
        maths.append(m.group(1).strip())
        return f"\x00MATH{len(maths) - 1}I\x00"

    md_text = _MATH_INLINE_RE.sub(inline_repl, md_text)
    return md_text, maths


def _extract_footnotes(md_text: str) -> tuple[str, list[tuple[int, str, str]]]:
    """收集 [^id]: 定义并删除；正文 [^id] 引用替换为上标占位符 \\x00FNn\\x00。

    返回 (处理后的文本, [(序号, id, 内容), ...])。引用按首次出现顺序编号。
    """
    defs: dict[str, str] = {}

    def def_repl(m):
        defs[m.group(1)] = m.group(2).strip()
        return ""  # 删除定义行

    md_text = _FN_DEF_RE.sub(def_repl, md_text)

    # 按首次出现顺序给引用编号
    order = list(dict.fromkeys(_FN_REF_RE.findall(md_text)))
    id2num = {fid: i + 1 for i, fid in enumerate(order)}

    def ref_repl(m):
        n = id2num.get(m.group(1))
        return f"\x00FN{n}\x00" if n else m.group(0)

    md_text = _FN_REF_RE.sub(ref_repl, md_text)

    fn_list = [(id2num[fid], fid, defs[fid]) for fid in order if fid in defs]
    return md_text, fn_list


def _make_inline_link(text: str, url: str, title: str | None) -> str:
    """构造 markdown 行内链接 [text](url "title")。"""
    if title:
        return f"[{text}]({url} \"{title}\")"
    return f"[{text}]({url})"


def _extract_ref_links(md_text: str) -> str:
    """收集 [id]: url 定义并删除；把引用式链接 [t][id]/[t][]/[id] 转成行内链接。

    转成行内格式后，逐行内联解析即可正确渲染（无需 markdown 库的跨行 reference 解析）。
    """
    defs: dict[str, tuple[str, str | None]] = {}

    def def_repl(m):
        defs[m.group(1)] = (m.group(2), m.group(3))
        return ""

    md_text = _REF_DEF_RE.sub(def_repl, md_text)

    def full_repl(m):
        text, key = m.group(1), m.group(2)
        if key in defs:
            return _make_inline_link(text, *defs[key])
        return m.group(0)

    md_text = re.sub(r"\[([^\]]+)\]\[([^\]]+)\]", full_repl, md_text)

    # collapsed: [text][]
    def collapsed_repl(m):
        key = m.group(1)
        if key in defs:
            return _make_inline_link(key, *defs[key])
        return m.group(0)

    md_text = re.sub(r"\[([^\]]+)\]\[\]", collapsed_repl, md_text)

    # shortcut: [key]（仅当 key 是已定义的引用 id；(?!\() 排除已是行内链接的情况）
    def shortcut_repl(m):
        key = m.group(1)
        if key in defs:
            return _make_inline_link(key, *defs[key])
        return m.group(0)

    md_text = re.sub(r"\[([^\]^]+)\](?!\()", shortcut_repl, md_text)
    return md_text


def _unescape_lt_gt_pipe(md_text: str) -> str:
    """markdown 库不转义 < > |（不在其可转义字符集），用占位符补转义。

    仅在围栏代码块外处理（代码块内容原样保留），避免误伤代码。
    占位符由 _InlineParser 在生成 run 时还原为 < > |。
    """
    parts = re.split(r"(```.*?```)", md_text, flags=re.DOTALL)
    for i, p in enumerate(parts):
        if p.startswith("```"):
            continue
        p = p.replace(r"\<", "\x00LT\x00")
        p = p.replace(r"\>", "\x00GT\x00")
        p = p.replace(r"\|", "\x00PIPE\x00")
        parts[i] = p
    return "".join(parts)


# ---------------------------------------------------------------------------
# 内联渲染：markdown 库 → HTML → runs
# ---------------------------------------------------------------------------
# 占位符识别（公式 / 脚注 / 软换行）
_PLACEHOLDER_RE = re.compile("\x00(MATH\\d+[BI]|FN\\d+|BR|LT|GT|PIPE)\x00")


def _apply_strikethrough(html_frag: str) -> str:
    """把 ~~x~~ 转 <del>x</del>，但跳过 <code>...</code> 内部（代码里的 ~~ 不动）。

    markdown 核心不处理删除线，故在库输出后补一道。
    """
    parts = re.split(r"(<code>.*?</code>)", html_frag, flags=re.DOTALL)
    for i, p in enumerate(parts):
        if p.startswith("<code>"):
            continue
        parts[i] = re.sub(r"~~(.+?)~~", r"<del>\1</del>", p)
    return "".join(parts)


def _md_inline_to_html(text: str) -> str:
    """把单行/已合并的内联 markdown 文本转为 HTML 内联片段。

    走 markdown 库（extra + nl2br），再补删除线。剥离外层 <p>。
    """
    if not text:
        return ""
    try:
        import markdown
        html = markdown.markdown(text, extensions=["extra", "nl2br"])
    except ImportError:
        # 无 markdown 库时退化为最小转义
        return _html_lib.escape(text)

    # 去除外层 <p>...</p>（可能有多个，逐个剥离）
    html = re.sub(r"^<p>(.*)</p>$", r"\1", html.strip(), flags=re.DOTALL)
    # 多段时用 <br> 连接（理论上单行内联不会多段）
    html = html.replace("</p>\n<p>", "<br/>")
    html = html.replace("</p><p>", "<br/>")
    html = _apply_strikethrough(html)
    return html


class _InlineParser(HTMLParser):
    """递归遍历内联 HTML，按样式栈在段落上生成 run（含链接/图片/换行/占位还原）。"""

    _TAG_STYLE = {
        "em": "italic", "i": "italic",
        "strong": "bold", "b": "bold",
        "del": "strike", "s": "strike", "strike": "strike",
        "code": "code",
        "sup": "sup",
        "sub": "sub",
    }

    def __init__(self, paragraph, font_name: str, font_size):
        super().__init__(convert_charrefs=True)
        self.para = paragraph
        self.font = font_name
        self.size = font_size
        self.stack: list[dict] = []

    def _merged(self) -> dict:
        m = {"bold": False, "italic": False, "strike": False,
             "code": False, "sup": False, "sub": False, "link": None}
        for s in self.stack:
            for k, v in s.items():
                m[k] = v
        return m

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        key = self._TAG_STYLE.get(tag)
        if key:
            self.stack.append({key: True})
        elif tag == "a":
            self.stack.append({"link": a.get("href", "")})
        elif tag == "br":
            self.para.add_run().add_break()
        elif tag == "img":
            self._add_img(a)

    def handle_startendtag(self, tag, attrs):
        # 自闭合 <br/> <img/>
        if tag in ("br", "img"):
            self.handle_starttag(tag, attrs)

    def handle_endtag(self, tag):
        key = self._TAG_STYLE.get(tag) or ("link" if tag == "a" else None)
        if not key:
            return
        for i in range(len(self.stack) - 1, -1, -1):
            if key in self.stack[i]:
                del self.stack[i]
                break

    def handle_data(self, data):
        if not data:
            return
        # 占位符切分：公式 / 脚注 / 软换行
        pos = 0
        for m in _PLACEHOLDER_RE.finditer(data):
            if m.start() > pos:
                self._emit_text(data[pos:m.start()])
            self._emit_placeholder(m.group(1))
            pos = m.end()
        if pos < len(data):
            self._emit_text(data[pos:])

    # ---- run 生成 ----
    def _emit_text(self, text: str):
        style = self._merged()
        if style["link"]:
            _add_hyperlink(self.para, text, style["link"], self.font, self.size)
            return
        run = self.para.add_run(text)
        _apply_inline_style(run, style, self.font, self.size)

    def _emit_placeholder(self, key: str):
        if key == "BR":
            self.para.add_run().add_break()
            return
        if key == "LT":
            self._emit_text("<")
            return
        if key == "GT":
            self._emit_text(">")
            return
        if key == "PIPE":
            self._emit_text("|")
            return
        m = re.match(r"MATH(\d+)([BI])", key)
        if m:
            idx = int(m.group(1))
            if 0 <= idx < len(_MATH_CACHE):
                self._emit_math(_MATH_CACHE[idx])
            return
        m = re.match(r"FN(\d+)", key)
        if m:
            run = self.para.add_run(m.group(1))
            run.font.superscript = True
            _set_run_font(run, self.font, self.size)

    def _emit_math(self, expr: str):
        """公式：等宽数学字体 + 斜体（简方案，保证不被破坏且可读）。"""
        from docx.shared import RGBColor
        run = self.para.add_run(expr)
        run.italic = True
        _set_run_font(run, "Cambria Math", self.size)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    def _add_img(self, attrs: dict):
        from docx.shared import Pt, Inches, RGBColor
        alt = attrs.get("alt", "图片")
        src = (attrs.get("src") or "").strip()
        img_path = _resolve_image_path(_CURRENT_INPUT_PATH, src)
        if img_path and img_path.exists():
            run = self.para.add_run()
            run.add_picture(str(img_path), width=Inches(2))
        else:
            run = self.para.add_run(f"[{alt}]")
            _set_run_font(run, self.font, self.size)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.italic = True


def _apply_inline_style(run, style: dict, font_name: str, font_size):
    """根据合并后的内联样式设置 run 字体与效果。"""
    from docx.shared import RGBColor
    if style["bold"]:
        run.bold = True
    if style["italic"]:
        run.italic = True
    if style["strike"]:
        run.font.strike = True
    if style["sup"]:
        run.font.superscript = True
    if style["sub"]:
        run.font.subscript = True
    if style["code"]:
        _set_run_font(run, "Consolas", _Pt(10) if font_size is None else _Pt(min(font_size.pt, 10)))
        _set_run_shading(run, "E8E8E8")
        return
    _set_run_font(run, font_name, font_size)


def _Pt(value):
    from docx.shared import Pt
    return Pt(value)


def _add_inline_runs(paragraph, text: str, base_font_name: str = "Microsoft YaHei",
                     base_font_size=None):
    """解析内联 Markdown 格式并添加 runs 到段落（走 markdown 库，支持嵌套/转义/链接）。"""
    if not text:
        return
    html_frag = _md_inline_to_html(text)
    if not html_frag:
        return
    parser = _InlineParser(paragraph, base_font_name, base_font_size)
    parser.feed(html_frag)
    parser.close()


# ---------------------------------------------------------------------------
# 代码高亮：Pygments token → run 颜色
# ---------------------------------------------------------------------------
_PYG_STYLE = None


def _get_pyg_style():
    """惰性加载 Pygments 浅底配色（default），返回 style.styles 字典。"""
    global _PYG_STYLE
    if _PYG_STYLE is not None:
        return _PYG_STYLE
    try:
        from pygments.styles import get_style_by_name
        _PYG_STYLE = get_style_by_name("default").styles
    except Exception:
        _PYG_STYLE = {}
    return _PYG_STYLE


def _token_color(ttype) -> tuple[bool, bool, str | None]:
    """返回。从最具体 token 向上查找。"""
    styles = _get_pyg_style()
    t = ttype
    while t is not None:
        rule = styles.get(t)
        if rule:
            bold = "bold" in rule
            italic = "italic" in rule
            m = re.search(r"#([0-9a-fA-F]{6})", rule)
            return bold, italic, (m.group(1) if m else None)
        t = t.parent
    return False, False, None


# ---------------------------------------------------------------------------
# 表格列对齐解析
# ---------------------------------------------------------------------------
def _parse_table_align(separator: str) -> list[str]:
    """解析表格分隔行，返回每列对齐方式：'left'/'center'/'right'。

    规则：:--: center / --: right / :-- 或 --- left。
    """
    cells = [c.strip() for c in separator.strip().strip("|").split("|")]
    aligns = []
    for c in cells:
        if not c:
            aligns.append("left")
            continue
        left = c.startswith(":")
        right = c.endswith(":")
        if left and right:
            aligns.append("center")
        elif right:
            aligns.append("right")
        else:
            aligns.append("left")
    return aligns


def _apply_quote_style(paragraph, depth: int):
    """引用块段落样式：左边框 + 浅底 + 缩进，随层级递增。"""
    if depth <= 0:
        return
    from docx.shared import Cm, Pt
    palette = ["4472C4", "70AD47", "ED7D31", "FFC000"]
    _add_left_border(paragraph, palette[(depth - 1) % len(palette)], 18)
    _set_paragraph_shading(paragraph, "F2F7FC")
    paragraph.paragraph_format.left_indent = Cm(1.0 + 0.4 * (depth - 1))
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# 主转换器
# ---------------------------------------------------------------------------
@register
class MdToDocx(BaseConverter):
    name = "md-docx"
    source_formats = ["md", "markdown"]
    target_formats = ["docx"]
    description = "Markdown 转 Word 文档 (Mermaid 自动转图片插入)"
    dependencies = ["python-docx"]  # import 名 docx

    # 标题字号映射：h1 → h6
    HEADING_SIZES = {
        1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11,
    }

    def convert(self, input_path: Path, output_path: Path, **options) -> ConvertResult:
        from docx import Document
        from docx.shared import Pt

        global _CURRENT_INPUT_PATH, _MATH_CACHE, _FN_CACHE
        _CURRENT_INPUT_PATH = input_path
        self._embedded_img_count = 0

        text = input_path.read_text(encoding="utf-8")

        # 预处理：公式 → 脚注 → 引用式链接（顺序敏感）
        text, _MATH_CACHE = _extract_math(text)
        text, _FN_CACHE = _extract_footnotes(text)
        text = _extract_ref_links(text)
        text = _unescape_lt_gt_pipe(text)

        doc = Document()

        # 基础样式：统一微软雅黑
        style = doc.styles["Normal"]
        _set_style_font(style, "Microsoft YaHei", Pt(11))
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.3

        for level in range(1, 7):
            heading_style = doc.styles[f"Heading {level}"]
            _set_style_font(heading_style, "Microsoft YaHei", Pt(self.HEADING_SIZES[level]))
            heading_style.font.bold = True

        # 预提取 mermaid 块并渲染为图片
        mermaid_pattern = r"```mermaid\s*\n(.*?)```"
        mermaid_blocks = re.findall(mermaid_pattern, text, re.DOTALL)
        mermaid_images = {}
        if mermaid_blocks:
            try:
                tmp_dir = Path(tempfile.mkdtemp(prefix="mermaid_"))
                images = _render_mermaid_images(mermaid_blocks, tmp_dir)
                for i, img in enumerate(images):
                    mermaid_images[i] = img
            except Exception:
                pass

        # 块级渲染（递归，引用块内部会递归调用）
        self._mermaid_images = mermaid_images
        self._mermaid_idx = 0
        lines = text.split("\n")
        self._render_blocks(doc, lines, quote_depth=0)

        # 文末脚注
        if _FN_CACHE:
            self._add_footnotes_section(doc)

        doc.save(str(output_path))

        msg = "Markdown 已转为 Word 文档"
        parts = []
        if getattr(self, "_embedded_img_count", 0):
            parts.append(f"{self._embedded_img_count} 张本地图片")
        img_count = len([v for v in mermaid_images.values() if v.exists()])
        if img_count:
            parts.append(f"{img_count} 个 Mermaid 图表")
        if _FN_CACHE:
            parts.append(f"{len(_FN_CACHE)} 个脚注")
        if parts:
            msg += "，" + "、".join(parts) + "已插入"
        return ConvertResult(True, output_path=output_path, message=msg)

    # ---- 块级解析（递归） ----
    def _render_blocks(self, doc, lines: list[str], *, quote_depth: int = 0):
        """递归块级渲染。quote_depth>0 表示处于引用块内，段落套引用样式。"""
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        ordered_counters: dict[int, int] = {}  # 缩进层级 -> 当前序号
        i = 0
        n = len(lines)

        while i < n:
            line = lines[i]

            # 非有序列表行重置有序编号计数器
            if not re.match(r"^(\s*)\d+\.\s+", line):
                ordered_counters.clear()

            # === 标题（引用内也支持，但降级用较小字号由 heading 样式决定）===
            m = re.match(r"^(#{1,6})\s+(.+)$", line)
            if m:
                level = len(m.group(1))
                heading = doc.add_heading(level=level)
                _add_inline_runs(heading, m.group(2).strip(), base_font_size=None)
                i += 1
                continue

            # === Mermaid 代码块 → 插入图片 ===
            if line.strip().startswith("```mermaid"):
                code_lines, i = self._collect_fenced(lines, i)
                if self._mermaid_idx in self._mermaid_images and \
                        self._mermaid_images[self._mermaid_idx].exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(self._mermaid_images[self._mermaid_idx]),
                                    width=Inches(6))
                else:
                    self._add_code_block(doc, code_lines, "mermaid")
                self._mermaid_idx += 1
                continue

            # === 普通代码块 ===
            if line.strip().startswith("```"):
                lang = line.strip()[3:].strip()
                code_lines, i = self._collect_fenced(lines, i)
                self._add_code_block(doc, code_lines, lang)
                continue

            # === 表格 ===
            if "|" in line and i + 1 < n and \
                    re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
                rows, aligns, i = self._collect_table(lines, i)
                if rows:
                    self._add_table(doc, rows, aligns)
                continue

            # === 引用块（剥一层 > 递归）===
            if line.lstrip().startswith(">"):
                block_lines = []
                while i < n and (lines[i].lstrip().startswith(">") or
                                 (lines[i].strip() == "" and block_lines and
                                  i + 1 < n and lines[i + 1].lstrip().startswith(">"))):
                    block_lines.append(lines[i])
                    i += 1
                stripped = [self._strip_one_quote(l) for l in block_lines]
                # 去掉首尾空行
                while stripped and not stripped[0].strip():
                    stripped.pop(0)
                while stripped and not stripped[-1].strip():
                    stripped.pop()
                if stripped:
                    self._render_blocks(doc, stripped, quote_depth=quote_depth + 1)
                continue

            # === 任务列表 ===
            m = re.match(r"^(\s*)[-*+]\s+\[([ xX])\]\s+(.+)$", line)
            if m:
                indent = _get_indent_level(line)
                checked = m.group(2).lower() == "x"
                content = m.group(3)
                prefix = "☑ " if checked else "☐ "
                # 普通段落（不用 List Bullet，避免叠加项目符号）
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.74 + 0.5 * indent)
                if quote_depth > 0:
                    _apply_quote_style(p, quote_depth)
                _add_inline_runs(p, prefix + content)
                i += 1
                continue

            # === 无序列表 ===
            m = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
            if m:
                indent = _get_indent_level(line)
                content = m.group(2)
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(1.27 * (indent + 1))
                if quote_depth > 0:
                    _apply_quote_style(p, quote_depth)
                _add_inline_runs(p, content)
                i += 1
                continue

            # === 有序列表（手动序号，章节重置）===
            m = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
            if m:
                indent = _get_indent_level(line)
                content = m.group(2)
                for k in [k for k in ordered_counters if k > indent]:
                    del ordered_counters[k]
                ordered_counters[indent] = ordered_counters.get(indent, 0) + 1
                num = ordered_counters[indent]
                # 普通段落（不用 List Number，避免全文连号）；前缀单独成 run，
                # 不经过 markdown（否则 "1. " 会被重新当成有序列表语法解析）
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0 + 0.5 * indent)
                p.paragraph_format.first_line_indent = Cm(-0.5)
                if quote_depth > 0:
                    _apply_quote_style(p, quote_depth)
                prefix_run = p.add_run(f"{num}. ")
                _set_run_font(prefix_run, "Microsoft YaHei", None)
                _add_inline_runs(p, content)
                i += 1
                continue

            # === 分割线 ===
            if re.match(r"^\s*(---+|\*\*\*+|___+)\s*$", line.strip()):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                self._add_hr_border(p)
                i += 1
                continue

            # === 空行 ===
            if not line.strip():
                i += 1
                continue

            # === 独立行图片 ===
            m = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
            if m:
                alt_text = m.group(1) or "图片"
                src = m.group(2).strip()
                self._add_standalone_image(doc, _CURRENT_INPUT_PATH, src, alt_text)
                i += 1
                continue

            # === 段落（合并连续普通行 + 软换行）===
            para_lines = []
            while i < n:
                ln = lines[i]
                if not ln.strip():
                    break
                if self._is_block_start(ln, lines, i):
                    break
                para_lines.append(ln)
                i += 1
            if para_lines:
                merged = self._merge_soft_breaks(para_lines)
                p = doc.add_paragraph()
                if quote_depth > 0:
                    _apply_quote_style(p, quote_depth)
                _add_inline_runs(p, merged)
            else:
                i += 1  # 防御：避免死循环

    # ---- 块级辅助 ----
    def _collect_fenced(self, lines, i):
        """收集围栏代码块内容，返回 (code_lines, next_i)。"""
        i += 1  # 跳过开围栏
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i])
            i += 1
        i += 1  # 跳过闭围栏
        return code_lines, i

    def _collect_table(self, lines, i):
        """收集表格行与列对齐，返回 (rows, aligns, next_i)。"""
        separator = lines[i + 1]
        aligns = _parse_table_align(separator)
        rows = []
        i += 0
        # 从当前行（表头）开始
        while i < len(lines) and "|" in lines[i] and lines[i].strip():
            raw = lines[i].strip()
            if re.match(r"^\|[\s\-:|]+\|$", raw):
                i += 1
                continue
            cells = [c.strip() for c in raw.strip("|").split("|")]
            rows.append(cells)
            i += 1
        return rows, aligns, i

    def _is_block_start(self, line: str, lines: list[str], i: int) -> bool:
        """判断 line 是否是块级元素起始（用于段落合并终止）。"""
        s = line.lstrip()
        if re.match(r"^#{1,6}\s+", s):
            return True
        if s.startswith("```"):
            return True
        if s.startswith(">"):
            return True
        if re.match(r"^(\s*)[-*+]\s+(\[[ xX]\])?", line):
            return True
        if re.match(r"^(\s*)\d+\.\s+", line):
            return True
        if re.match(r"^\s*(---+|\*\*\*+|___+)\s*$", s):
            return True
        if re.match(r"^\s*!\[[^\]]*\]\([^)]+\)\s*$", s):
            return True
        # 表格：本行有 | 且下一行是分隔行
        if "|" in line and i + 1 < len(lines) and \
                re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            return True
        return False

    def _strip_one_quote(self, line: str) -> str:
        """剥掉引用行最前面的一个 > 及可选空格（>> x → > x；> x → x）。"""
        return re.sub(r"^\s*>\s?", "", line, count=1)

    def _merge_soft_breaks(self, para_lines: list[str]) -> str:
        """合并段落行：行尾有空格 → 软换行占位，否则空格连接。"""
        out = ""
        for ln in para_lines:
            if ln.rstrip() != ln:
                # 行尾有空格/tab → 软换行（markdown 两空格 + 回车）
                out += ln.rstrip() + "\x00BR\x00"
            else:
                out += ln + " "
        return out.strip()

    # ---- 复合元素 ----
    def _add_standalone_image(self, doc, input_path: Path, src: str, alt_text: str):
        """插入单独成行的图片，读取本地文件并居中嵌入。"""
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        img_path = _resolve_image_path(input_path, src)
        if img_path and img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(6))
            self._embedded_img_count += 1
            if alt_text and alt_text != "图片":
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(alt_text)
                _set_run_font(cap_run, "Microsoft YaHei", Pt(9))
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                cap_run.italic = True
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[图片缺失: {alt_text}]")
            _set_run_font(run, "Microsoft YaHei", Pt(10))
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True

    def _add_code_block(self, doc, code_lines: list[str], lang: str = ""):
        """添加代码块：单段落 + Pygments token 上色 + 浅灰底。"""
        from docx.shared import Pt, Cm, RGBColor

        code = "\n".join(code_lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.15
        _set_paragraph_shading(p, "F5F5F5")

        tokens = self._lex_code(code, lang)
        if not tokens:
            run = p.add_run(" ")
            _set_run_font(run, "Consolas", Pt(9.5))
            return

        for ttype, value in tokens:
            bold, italic, hexcolor = _token_color(ttype)
            parts = value.split("\n")
            for k, part in enumerate(parts):
                if k > 0:
                    p.add_run().add_break()  # 段内换行，保持同段底色
                if not part:
                    continue
                run = p.add_run(part)
                _set_run_font(run, "Consolas", Pt(9.5))
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if hexcolor:
                    run.font.color.rgb = RGBColor(int(hexcolor[0:2], 16),
                                                   int(hexcolor[2:4], 16),
                                                   int(hexcolor[4:6], 16))

    def _lex_code(self, code: str, lang: str):
        """用 Pygments 词法分析，返回 token 列表；失败则返回整体纯文本。"""
        if not code:
            return []
        try:
            from pygments import lex
            from pygments.lexers import get_lexer_by_name, TextLexer
            try:
                lexer = get_lexer_by_name(lang) if lang else TextLexer()
            except Exception:
                lexer = TextLexer()
            return list(lex(code, lexer))
        except ImportError:
            return [(None, code)]

    def _add_table(self, doc, rows: list[list[str]], aligns: list[str] | None = None):
        """添加表格：首行加粗，按列对齐。"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not rows:
            return
        aligns = aligns or []
        max_cols = max(len(r) for r in rows)
        # 对齐不足时补 left
        aligns = (aligns + ["left"] * max_cols)[:max_cols]
        align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                     "center": WD_ALIGN_PARAGRAPH.CENTER,
                     "right": WD_ALIGN_PARAGRAPH.RIGHT}

        table = doc.add_table(rows=len(rows), cols=max_cols, style="Table Grid")
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx >= max_cols:
                    break
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                para = cell.paragraphs[0]
                para.alignment = align_map.get(aligns[c_idx], WD_ALIGN_PARAGRAPH.LEFT)
                _add_inline_runs(para, cell_text)
                if r_idx == 0:
                    for run in para.runs:
                        run.bold = True

    def _add_footnotes_section(self, doc):
        """文末追加脚注列表。"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        if not _FN_CACHE:
            return
        # 分割线
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(12)
        sep.paragraph_format.space_after = Pt(4)
        self._add_hr_border(sep)
        title = doc.add_paragraph()
        title_run = title.add_run("脚注")
        _set_run_font(title_run, "Microsoft YaHei", Pt(11))
        title_run.bold = True

        for num, _fid, content in _FN_CACHE:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            p.paragraph_format.first_line_indent = Pt(-12)
            _add_inline_runs(p, f"{num}. {content}", base_font_size=Pt(10))

    def _add_hr_border(self, paragraph):
        """用底部边框模拟水平分割线。"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)
