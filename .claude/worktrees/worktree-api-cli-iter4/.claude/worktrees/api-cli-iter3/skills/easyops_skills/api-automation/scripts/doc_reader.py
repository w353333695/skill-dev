#!/usr/bin/env python3
"""
文档读取工具 - 支持 Word (.docx) 和 PDF 文件
用于从文档中提取 API 相关内容
"""

import argparse
import sys
import json


def read_docx(file_path: str) -> str:
    """读取 Word 文档内容"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装 python-docx 库", file=sys.stderr)
        print("运行: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(file_path)
    content = []

    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)

    # 读取表格内容
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(" | ".join(row_data))
        if table_data:
            content.append("\n[表格]\n" + "\n".join(table_data))

    return "\n\n".join(content)


def read_pdf(file_path: str) -> str:
    """读取 PDF 文档内容"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("错误: 需要安装 PyMuPDF 库", file=sys.stderr)
        print("运行: python3 -m pip install PyMuPDF", file=sys.stderr)
        sys.exit(1)

    doc = fitz.open(file_path)
    content = []

    for page_num, page in enumerate(doc, 1):
        text = page.get_text()
        if text.strip():
            content.append(f"[第 {page_num} 页]\n{text}")

    doc.close()
    return "\n\n".join(content)


def main():
    parser = argparse.ArgumentParser(description="文档读取工具")
    parser.add_argument("--type", "-t", choices=["docx", "pdf"], required=True,
                        help="文档类型: docx 或 pdf")
    parser.add_argument("--file", "-f", required=True, help="文档路径")
    parser.add_argument("--output", "-o", choices=["text", "json"], default="text",
                        help="输出格式: text 或 json")

    args = parser.parse_args()

    if args.type == "docx":
        content = read_docx(args.file)
    else:
        content = read_pdf(args.file)

    if args.output == "json":
        print(json.dumps({"file": args.file, "type": args.type, "content": content},
                         ensure_ascii=False, indent=2))
    else:
        print(content)


if __name__ == "__main__":
    main()
